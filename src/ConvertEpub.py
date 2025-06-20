import os
import re
import subprocess 
from docx import Document 
from docx.enum.section import WD_SECTION 
# from docx.shared import Inches # Đã bỏ comment vì không sử dụng

# Định nghĩa các hằng số để tránh lỗi f-string với backslash
NEWLINE_CHAR = "\n"
BREAK_TAG = "<br/>"

# RẤT QUAN TRỌNG:
# HÃY THAY THẾ DÒNG DƯỚI ĐÂY BẰNG ĐƯỜNG DẪN THẬT CỦA pandoc.exe TRÊN MÁY BẠN.
# Bạn có thể tìm thấy nó bằng cách mở Command Prompt/PowerShell và gõ: where pandoc
# Ví dụ: PANDOC_PATH = r"C:\Program Files\Pandoc\pandoc.exe"
# Ví dụ khác: PANDOC_PATH = r"C:\Users\YOUR_USERNAME\AppData\Local\Pandoc\pandoc.exe"
PANDOC_PATH = r"C:\Users\vinhd\AppData\Local\Pandoc\pandoc.exe" # <--- CHỈNH SỬA DÒNG NÀY!


def txt_to_docx(txt_path, docx_path, book_title, chapter_pattern):
    """
    Chuyển đổi file .txt sang .docx, tự động nhận diện chương và tạo header
    với page break.

    Args:
        txt_path (str): Đường dẫn đến file .txt nguồn.
        docx_path (str): Đường dẫn đến file .docx đích.
        book_title (str): Tiêu đề của sách (sẽ dùng cho tiêu đề H1 nếu không có intro).
        chapter_pattern (str): Biểu thức chính quy (regex) để nhận diện các chương.
    Returns:
        bool: True nếu thành công, False nếu thất bại.
    """
    print(f"Bắt đầu giai đoạn 1/2: Chuyển đổi TXT sang DOCX tại '{docx_path}'...")

    if not os.path.exists(txt_path):
        print(f"  Lỗi: Không tìm thấy file .txt tại đường dẫn '{txt_path}'")
        return False

    # 1. Đọc nội dung file .txt
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            full_text = f.read()
        print("  Đã đọc xong file .txt.")
    except Exception as e:
        print(f"  Lỗi khi đọc file .txt: {e}")
        return False

    # 2. Khởi tạo tài liệu DOCX
    document = Document()
    # document.sections[0].left_margin = Inches(1) # Ví dụ: thiết lập lề
    # document.sections[0].right_margin = Inches(1)
    # document.sections[0].top_margin = Inches(1)
    # document.sections[0].bottom_margin = Inches(1)

    # 3. Tìm kiếm các chương bằng regex
    matches = list(re.finditer(chapter_pattern, full_text, re.MULTILINE))
    total_chapters = len(matches)

    if total_chapters == 0 and full_text.strip():
        print("  Không tìm thấy định dạng chương nào. Toàn bộ file sẽ được coi là một trang DOCX duy nhất.")
        document.add_heading(book_title, level=1)
        paragraphs = full_text.strip().split('\n\n')
        for para in paragraphs:
            if para.strip():
                document.add_paragraph(para.replace(NEWLINE_CHAR, ' ')) 

    else:
        print(f"  Đã tìm thấy {total_chapters} chương. Đang tạo nội dung DOCX...")
        
        intro_text_handled = False
        # Xử lý phần giới thiệu (nếu có) trước chương đầu tiên
        if matches and matches[0].start() > 0:
            intro_text = full_text[0:matches[0].start()].strip()
            if intro_text:
                print("  Đang xử lý phần Mở Đầu...")
                document.add_heading('Mở Đầu', level=1)
                paragraphs = intro_text.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        document.add_paragraph(para.replace(NEWLINE_CHAR, ' '))
                # Thêm page break sau phần mở đầu
                document.add_page_break()
                print("  Đã xử lý xong phần Mở Đầu.")
                intro_text_handled = True

        # Xử lý từng chương
        for i, match in enumerate(matches):
            chapter_title = match.group(0).strip()
            start_index = match.end()

            if i + 1 < len(matches):
                end_index = matches[i+1].start()
            else:
                end_index = len(full_text)

            chapter_content = full_text[start_index:end_index].strip()

            # Thêm page break trước mỗi chương (trừ chương đầu tiên nếu nó là chương đầu tiên của tài liệu
            # và không có phần giới thiệu đi trước nó)
            if i > 0 or intro_text_handled: 
                document.add_page_break()

            # Thêm tiêu đề chương (Heading 1)
            document.add_heading(chapter_title, level=1)
            
            # Thêm nội dung chương
            if chapter_content:
                paragraphs = chapter_content.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        document.add_paragraph(para.replace(NEWLINE_CHAR, ' ')) 

            print(f"\r  - Đã xử lý chương {i+1}/{total_chapters}: {chapter_title[:50]}...", end='')
        print("\n  Đã xử lý xong tất cả các chương.")

    # 4. Ghi file DOCX
    try:
        document.save(docx_path)
        print(f"  Thành công! File DOCX đã được tạo tại: {docx_path}")
        return True
    except Exception as e:
        print(f"  Lỗi khi ghi file DOCX: {e}")
        return False

def docx_to_epub(docx_path, epub_path, book_title, book_author):
    """
    Chuyển đổi file .docx sang .epub bằng Pandoc.

    Args:
        docx_path (str): Đường dẫn đến file .docx nguồn.
        epub_path (str): Đường dẫn đến file .epub đích.
        book_title (str): Tiêu đề của sách.
        book_author (str): Tác giả của sách.
    Returns:
        bool: True nếu thành công, False nếu thất bại.
    """
    print(f"\nBắt đầu giai đoạn 2/2: Chuyển đổi DOCX sang EPUB tại '{epub_path}' bằng Pandoc...")

    if not os.path.exists(docx_path):
        print(f"  Lỗi: Không tìm thấy file .docx tại đường dẫn '{docx_path}'")
        return False

    # Xây dựng lệnh Pandoc
    command = [
        PANDOC_PATH, # <--- Đã thay đổi từ 'pandoc' thành PANDOC_PATH
        '-o', epub_path,
        docx_path,
        '--metadata', f'title={book_title}',
        '--metadata', f'author={book_author}',
        '--epub-chapter-level=1', 
        '--toc', 
        '--toc-depth=1' 
    ]

    try:
        # Chạy lệnh Pandoc
        result = subprocess.run(command, check=True, capture_output=True, text=True, encoding='utf-8')
        
        print(f"  Thành công! File EPUB đã được tạo tại: {epub_path}")
        if result.stdout:
            print("  Pandoc output (stdout):")
            print(result.stdout)
        if result.stderr:
            print("  Pandoc errors/warnings (stderr):")
            print(result.stderr)
        return True
    except FileNotFoundError:
        print("\n  Lỗi: Pandoc không được tìm thấy TẠI ĐƯỜNG DẪN ĐÃ CHỈ ĐỊNH.")
        print(f"  Đảm bảo đường dẫn '{PANDOC_PATH}' là chính xác.")
        print("  Sử dụng 'where pandoc' trong CMD/PowerShell để tìm đường dẫn của bạn.")
        print("  Tải Pandoc tại: https://pandoc.org/installing.html")
        return False
    except subprocess.CalledProcessError as e:
        print(f"\n  Lỗi khi chạy Pandoc. Mã lỗi: {e.returncode}")
        print(f"  Pandoc stdout: {e.stdout}")
        print(f"  Pandoc stderr: {e.stderr}")
        print(f"  Lệnh đã chạy: {' '.join(command)}")
        return False
    except Exception as e:
        print(f"  Lỗi không xác định khi chuyển đổi DOCX sang EPUB: {e}")
        return False

# --- Phần chạy chính của script ---
if __name__ == "__main__":
    print("\n--- Chuyển đổi TXT sang DOCX sang EPUB ---")

    txt_input = input("Nhập đường dẫn đến file .txt nguồn (ví dụ: my_story.txt): ")
    txt_file_path = os.path.abspath(txt_input) 

    # Tự động đặt tên file DOCX và EPUB theo tên file TXT
    # Ví dụ: my_story.txt -> my_story.docx và my_story.epub (trong cùng thư mục)
    txt_base_name = os.path.splitext(os.path.basename(txt_file_path))[0]
    
    docx_file_path = os.path.join(os.path.dirname(txt_file_path), txt_base_name + ".docx")
    epub_file_path = os.path.join(os.path.dirname(txt_file_path), txt_base_name + ".epub")

    title = input(f"Nhập tiêu đề sách (mặc định: '{txt_base_name}'): ")
    if not title:
        title = txt_base_name
        print(f"Sử dụng tiêu đề mặc định: {title}")

    author = input("Nhập tên tác giả (mặc định: Unknown Author): ")
    if not author:
        author = "Unknown Author"
        print(f"Sử dụng tác giả mặc định: {author}")

    default_chapter_regex = r"^Chương\s+\d+:\s+.*$"
    
    custom_regex = input(
        f"Nhập biểu thức chính quy (regex) để nhận diện chương (mặc định: '{default_chapter_regex}'):\n"
        f"  (Ví dụ: '^Phần\\s+\\d+\\.\\s+.*$' cho 'Phần 1. Tên Phần')\n"
        f"Nhập bỏ trống để dùng mặc định: "
    )

    final_chapter_regex = custom_regex if custom_regex else default_chapter_regex

    print("\n--- Thông tin chuyển đổi ---")
    print(f"  File TXT nguồn: {txt_file_path}")
    print(f"  File DOCX trung gian: {docx_file_path}")
    print(f"  File EPUB đích: {epub_file_path}") # Giờ đây sẽ hiển thị đầy đủ tên file
    print(f"  Tiêu đề sách: {title}")
    print(f"  Tác giả: {author}")
    print(f"  Regex nhận diện chương: '{final_chapter_regex}'")
    print("--------------------------")

    success_docx = False
    try:
        success_docx = txt_to_docx(txt_file_path, docx_file_path, title, final_chapter_regex)
    except Exception as e:
        print(f"\nLỗi nghiêm trọng trong quá trình tạo DOCX: {e}")

    if success_docx:
        try:
            docx_to_epub(docx_file_path, epub_file_path, title, author)
        except Exception as e:
            print(f"\nLỗi nghiêm trọng trong quá trình tạo EPUB: {e}")
    else:
        print("\nSkipping EPUB conversion due to DOCX creation failure.")

    print("\nQuá trình chuyển đổi đã hoàn tất.")